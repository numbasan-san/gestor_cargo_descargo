
# Import docx NOT python-docx
import docx
 
# Create an instance of a word document
doc = docx.Document()
 
# Choosing the top most section of the page
section = doc.sections[0]
 
# Selecting the header
header = section.header
 
# Selecting the paragraph already present in
# the header section
header_para = header.paragraphs[0]
 
# Adding the centred zoned header
header_para.text = "\tThis is Centred Zoned Header..."
 
# Add a Title to the document
doc.add_heading('GeeksForGeeks', 0)
 
# Now save the document to a location
doc.save('gfg.docx')



# -*- coding: utf-8 -*-

'''
Implement floating image based on python-docx 0.8.10.

- Text wrapping style: BEHIND TEXT <wp:anchor behindDoc="1">
- Picture position: top-left corner of PAGE `<wp:positionH relativeFrom="page">`.

Create a docx sample (Layout | Positions | More Layout Options) and explore the 
source xml (Open as a zip | word | document.xml) to implement other text wrapping
styles and position modes per `CT_Anchor._anchor_xml()`.
'''

from docx.oxml import parse_xml, register_element_cls
from docx.oxml.ns import nsdecls
from docx.oxml.shape import CT_Picture
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne

# refer to docx.oxml.shape.CT_Inline
class CT_Anchor(BaseOxmlElement):
    """
    ``<w:anchor>`` element, container for a floating image.
    """
    extent = OneAndOnlyOne('wp:extent')
    docPr = OneAndOnlyOne('wp:docPr')
    graphic = OneAndOnlyOne('a:graphic')

    @classmethod
    def new(cls, cx, cy, shape_id, pic, pos_x, pos_y):
        """
        Return a new ``<wp:anchor>`` element populated with the values passed
        as parameters.
        """
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = shape_id
        anchor.docPr.name = 'Picture %d' % shape_id
        anchor.graphic.graphicData.uri = (
            'http://schemas.openxmlformats.org/drawingml/2006/picture'
        )
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def new_pic_anchor(cls, shape_id, rId, filename, cx, cy, pos_x, pos_y):
        """
        Return a new `wp:anchor` element containing the `pic:pic` element
        specified by the argument values.
        """
        pic_id = 0  # Word doesn't seem to use this, but does not omit it
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        anchor = cls.new(cx, cy, shape_id, pic, pos_x, pos_y)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def _anchor_xml(cls, pos_x, pos_y):
        return (
            '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" \n'
            '           behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" \n'
            '           %s>\n'
            '  <wp:simplePos x="0" y="0"/>\n'
            '  <wp:positionH relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionH>\n'
            '  <wp:positionV relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionV>\n'                    
            '  <wp:extent cx="914400" cy="914400"/>\n'
            '  <wp:wrapNone/>\n'
            '  <wp:docPr id="666" name="unnamed"/>\n'
            '  <wp:cNvGraphicFramePr>\n'
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            '  </wp:cNvGraphicFramePr>\n'
            '  <a:graphic>\n'
            '    <a:graphicData uri="URI not set"/>\n'
            '  </a:graphic>\n'
            '</wp:anchor>' % ( nsdecls('wp', 'a', 'pic', 'r'), int(pos_x), int(pos_y) )
        )


# refer to docx.parts.story.BaseStoryPart.new_pic_inline
def new_pic_anchor(part, image_descriptor, width, height, pos_x, pos_y):
    """Return a newly-created `w:anchor` element.

    The element contains the image specified by *image_descriptor* and is scaled
    based on the values of *width* and *height*.
    """
    rId, image = part.get_or_add_image(image_descriptor)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, filename = part.next_id, image.filename    
    return CT_Anchor.new_pic_anchor(shape_id, rId, filename, cx, cy, pos_x, pos_y)


# refer to docx.text.run.add_picture
def add_float_picture(p, image_path_or_stream, width=None, height=None, pos_x=0, pos_y=0):
    """Add float picture at fixed position `pos_x` and `pos_y` to the top-left point of page.
    """
    run = p.add_run()
    anchor = new_pic_anchor(run.part, image_path_or_stream, width, height, pos_x, pos_y)
    run._r.add_drawing(anchor)

# refer to docx.oxml.shape.__init__.py
register_element_cls('wp:anchor', CT_Anchor)


# if __name__ == '__main__':

from docx import Document
from docx.shared import Inches, Pt

document = Document()

# add a floating image
p = document.add_paragraph()
add_float_picture(p, 'logo.jfif', width=Pt(100), pos_x=Pt(500), pos_y=Pt(30))

# add text
p.add_run('Hello World'*50)


document.save('output.docx')


'''
import docx
from docx import *
from docx.shared import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = docx.Document('plantilla.docx')
# doc = docx.Document()
def create_paragraph(paragraph, alignment, text, style, size, bold = False):
    # paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  if alignment == 'justify' else WD_ALIGN_PARAGRAPH.CENTER if alignment == 'center' else WD_ALIGN_PARAGRAPH.RIGHT if alignment == 'right' else WD_ALIGN_PARAGRAPH.LEFT # WD_ALIGN_PARAGRAPH.LEFT
    txt = paragraph.add_run(text)
    txt.bold = bold
    font = txt.font
    font.name = style
    font.size = Pt(size)

def create_footer(text, style, size, bold = False):
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0].add_run(text)
    paragraph.bold = bold
    stylo = paragraph.style
    stylo.font.name = style
    stylo.font.size = Pt(size)

def create_header(text, style, size, bold = False):
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0].add_run(text + '\t')
    paragraph.add_picture('logo.jfif', width=Pt(100), height=Pt(100))
    paragraph.bold = bold
    stylo = paragraph.style
    stylo.font.name = style
    stylo.font.size = Pt(size)

def jump_line():
    paragraph = doc.add_paragraph().add_run().add_break()

def create_table(data):
    table = doc.add_table(rows=1, cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    color_table(table, 0, 0, '6EADFF')
    color_table(table, 0, 1, '6EADFF')

    table.rows[0].cells[0].text = 'DESCRIPCIÓN'
    table.rows[0].cells[0].text.format
    table.rows[0].cells[1].text = 'DETALLES'

    i = 1
    for title, content in data:
        row_cells = table.add_row().cells

        color_table(table, i, 0, 'CCD9FF')

        row_cells[0].text = title
        row_cells[1].text = str(content)
        table.cell(i, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        # make_rows_bold(table.rows[0], table.columns[0], bold = True, style = 'Calibri', size = 12)
        make_rows_bold(table.rows[1], table.columns[1], bold = False, style = 'Calibri', size = 12)
        make_rows_bold(table.rows[0], table.columns[0], bold = True, style = 'Calibri', size = 12)
        
        
        i += 1

    table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    table.columns[0].width = Cm(2)
    for cell in table.column_cells(0):
        cell.width = Cm(5)
    table.columns[1].width = Cm(2)
    for cell in table.column_cells(1):
        cell.width = Cm(7)

def color_table(table, row, cell, color):
    cll_xml_elm = table.rows[row].cells[cell]._tc
    tbl_cll_properties = cll_xml_elm.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), color)
    tbl_cll_properties.append(shade_obj)

def make_rows_bold(*rows, bold = False, style, size):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = bold
                    run.font.name = style
                    run.font.size = Pt(size)

nombre = 'qwerty uiop'
tipo_registro = 'cargo'
tipo_equipo = 'laptop'
dia = 3
mes = 5
año = 2021
equipos = 'mac'
seriales = '12345'

def getText():
    for para in doc.paragraphs:
        if '${tipo_registro}' in para.text:
            text = (para.text).replace('${tipo_registro}', tipo_registro)
            para.text = ''
            title = 'ANEXO 02 - CARGO' if tipo_registro == 'cargo' else 'ANEXO 04 - DESCARGO'
            create_paragraph(para, 'center', f'{title} EQUIPOS ELECTRÓNICOS', 'Calibri', 21, bold = True)
        if '${nombre}' in para.text:
            text = (para.text).replace('${nombre}', nombre)
            para.text = ''
            # para.add_run(text)
            create_paragraph(para, 'justify', text, 'Calibri', 12)
        if '${tipo_equipo}' in para.text:
            text = (para.text).replace('${tipo_equipo}', tipo_equipo)
            para.text = ''
            # para.add_run(text)
            create_paragraph(para, 'justify', text, 'Calibri', 12)

    doc.save('new.docx')

print(getText())
'''
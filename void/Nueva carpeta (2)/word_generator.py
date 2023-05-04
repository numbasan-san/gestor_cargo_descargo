
from docx import *
from docx.shared import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = Document()
def create_paragraph(alignment, text, style, size, bold = False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  if alignment == 'justify' else WD_ALIGN_PARAGRAPH.CENTER if alignment == 'center' else WD_ALIGN_PARAGRAPH.RIGHT if alignment == 'right' else WD_ALIGN_PARAGRAPH.LEFT # WD_ALIGN_PARAGRAPH.LEFT
    txt = paragraph.add_run(text)
    txt.bold = bold
    font = txt.font
    font.name = style
    font.size = Pt(size)

def create_footer(text, style, size, bold = False):
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0].add_run(text)
    paragraph.bold = bold
    stylo = paragraph.style
    stylo.font.name = style
    stylo.font.size = Pt(size)

def create_header(text, style, size, bold = False):
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0].add_run(text + '\t')
    paragraph.add_picture('logo.jfif', width=Pt(100), height=Pt(100))
    paragraph.bold = bold
    stylo = paragraph.style
    stylo.font.name = style
    stylo.font.size = Pt(size)

    
    ''' htable = header.add_table(1, 2, Inches(6))
htab_cells = htable.rows[0].cells
ht0=htab_cells[0].add_paragraph()
kh=ht0.add_run()
kh.add_picture('logo.jfif', width=Inches(1))'''
    
    """pic = doc.sections[0]
    pic.add_picture('logo.jfif', width=Pt(100), height=Pt(100))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT"""

def jump_line():
    paragraph = doc.add_paragraph().add_run().add_break()

def create_table(data):
    table = doc.add_table(rows=1, cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    color_table(table, 0, 0, '6EADFF')
    color_table(table, 0, 1, '6EADFF')

    table.rows[0].cells[0].text = 'DESCRIPCIÓN'
    table.rows[0].cells[0].text.format
    table.rows[0].cells[1].text = 'DETALLES'

    i = 1
    for title, content in data:
        row_cells = table.add_row().cells

        color_table(table, i, 0, 'CCD9FF')

        row_cells[0].text = title
        row_cells[1].text = str(content)
        table.cell(i, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        # make_rows_bold(table.rows[0], table.columns[0], bold = True, style = 'Calibri', size = 12)
        make_rows_bold(table.rows[1], table.columns[1], bold = False, style = 'Calibri', size = 12)
        make_rows_bold(table.rows[0], table.columns[0], bold = True, style = 'Calibri', size = 12)
        
        
        i += 1

    table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    table.columns[0].width = Cm(2)
    for cell in table.column_cells(0):
        cell.width = Cm(5)
    table.columns[1].width = Cm(2)
    for cell in table.column_cells(1):
        cell.width = Cm(7)

def color_table(table, row, cell, color):
    cll_xml_elm = table.rows[row].cells[cell]._tc
    tbl_cll_properties = cll_xml_elm.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), color)
    tbl_cll_properties.append(shade_obj)

def make_rows_bold(*rows, bold = False, style, size):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = bold
                    run.font.name = style
                    run.font.size = Pt(size)

# document.add_heading('${tipo_registro} EQUIPOS LECTRÓNICOS', 0)
nombre = 'qwerty uiop'
tipo_registro = 'cargo'
tipo_equipo = 'laptop'
dia = 3
mes = 5
año = 2021
equipos = 'mac'
seriales = '12345'

create_header(f'\tUso Interno', 'Calibri', 12, bold = True)

"""
paragraph = doc.add_paragraph()
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_picture('logo.jfif', width=Pt(100), height=Pt(100))
last_paragraph = doc.paragraphs[0] 
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
paragraph_containing_picture = doc.paragraphs[-1]
paragraph_containing_picture.style = 'centerstyle'
"""
"""
"""
# Título del documento.
title = 'ANEXO 02 - CARGO' if tipo_registro == 'cargo' else 'ANEXO 04 - DESCARGO'
create_paragraph('center', f'{title} EQUIPOS ELECTRÓNICOS', 'Calibri', 21, bold = True)

# Salto de línea.
# jump_line()

# Información del documento.
text = f'''Quien suscribe, La Sr(a).: {nombre}, Dominicano, mayor de edad, portador de la cédula de identidad y electoral No.: ____________________________________, domiciliado y residente en , ____________________________________ de esta Ciudad de Santo Domingo, por medio del presente documento DECLARO, haberle entregado el/al Sr(a).: ____________________________________, portador de la cédula de identidad y electoral No.: ____________________________________ , quien actúa como representante de la empresa GCS Systems, lo siguiente:'''
create_paragraph('justify', text, 'Calibri', 12)

# Salto de línea.
# paragraph = doc.add_paragraph().add_run().add_break()

# Tabla
data = (
    ('EQUIPO', tipo_equipo), 
    ('TIPO DE EQUIPO', tipo_equipo), 
    ('MARCA', equipos), 
    ('SERIAL', seriales), 
    ('NÚMERO TELEFÓNICO', '')
)
create_table(data)

text = f'''
Propiedad de la empresa GCS Systems LTD, los suscritos en el presente Recibo de Descargo, firman al pie del presente documento, en señal de aprobación de este.'''
create_paragraph('justify', text, 'Calibri', 12)

text = (f'''Al firmar este documento me comprometo a cumplir con los lineamientos establecidos en la POLÍTICA SERVICIOS TECNOLÓGICOS RCS-SEG-03-PL; así como con todas las políticas de Seguridad de la información y ciberseguridad de GCS Systems con sus procesos internos establecidos y divulgados.''')
create_paragraph('justify', text, 'Calibri', 12)

count = 'uno (1) original' if tipo_registro == 'cargo' else 'dos (2) originales'
text = (f'''Hecho en {count}, una para cada parte firmante, en la Ciudad de Santo Domingo, Distrito Nacional, Capital de la Republica dominicana, a los {dia} días del mes de {mes} del año {año}.''')
create_paragraph('justify', text, 'Calibri', 12)

text = '''_________________________________ _________________________________
Nombre y Firma			Nombre y Firma
    DECLARANTE		             REPRESENTANTE

'''
create_paragraph('center', text, 'Calibri', 12, bold = True)

jump_line()

text = """
RCS-SEG-03-PL\tUso Interno\tV3"""
create_footer(text, 'Calibri', 12, bold = True)

doc.save('output.docx')

"""
Lo mejor es declarar el párrafo, su organización, el texto, su fuente y el tamaño de su fuente; 
en ese orden.
"""

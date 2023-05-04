import sys
from docx import Document
from docx.shared import Inches

target_document = Document()
"""
target_document.sections[0].left_margin = Inches(0.3)
target_document.add_picture('logo.jfif', width=Inches(8.0))

new_section = target_document.add_section()
new_section.left_margin = Inches(1.0)
"""
source_document = Document('existing.docx')
print(source_document)
for paragraph in source_document.paragraphs:
     if "$nombre" in paragraph.text:
          new_txt = (paragraph.text).replace('$nombre', '0 1 1 2 3 5 8 13 21')
          paragraph.text = new_txt 
target_document.save('new.docx')

# logo.jfif
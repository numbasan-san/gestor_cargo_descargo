from docx import Document
from docx.shared import Inches

target_document = Document()
target_document.sections[0].left_margin = Inches(0.3)
target_document.add_picture('logo.jfif', width=Inches(8.0))

new_section = target_document.add_section()
new_section.left_margin = Inches(1.0)

source_document = Document('existing.docx')
for paragraph in source_document.paragraphs:
     target_document.add_paragraph(paragraph.text)
target_document.save('new.docx')

# logo.jfif